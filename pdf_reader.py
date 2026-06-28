import os
import docx
from pypdf import PdfReader

def extract_text_from_pdf(file_path_or_bytes) -> str:
    """
    Extracts text content from a PDF file path or file-like bytes object.
    """
    try:
        reader = PdfReader(file_path_or_bytes)
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to read PDF: {str(e)}")

def extract_text_from_docx(file_path_or_bytes) -> str:
    """
    Extracts text content from a DOCX file path or file-like bytes object.
    """
    try:
        doc = docx.Document(file_path_or_bytes)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text).strip()
    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {str(e)}")

def extract_text(file_source, filename: str) -> str:
    """
    Extracts text based on the file extension.
    Supports .pdf and .docx files.
    """
    ext = os.path.splitext(filename.lower())[1]
    if ext == ".pdf":
        return extract_text_from_pdf(file_source)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_source)
    else:
        # Try reading as plain text if it's not a known binary format
        try:
            if hasattr(file_source, "read"):
                # If bytes
                content = file_source.read()
                if isinstance(content, bytes):
                    return content.decode("utf-8", errors="ignore")
                return str(content)
            else:
                with open(file_source, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
        except Exception as e:
            raise ValueError(f"Unsupported file format or reading error: {ext}. {str(e)}")
